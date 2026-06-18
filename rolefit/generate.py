"""Stage-2 generation — tailored CV (.docx), cover letter, interview prep, learning
plan for a (person, job) match. Gated: meant to run on strong matches (≥ threshold)
or on demand. Artifacts saved under ~/.hermes/rolefit_artifacts/<slug>/<job_id>/.
"""

from __future__ import annotations

import json
import time
from pathlib import Path
from typing import Any, Optional

from . import db as _db
from . import jobs as _jobs
from . import llm as _llm
from . import scoring as _scoring


def _art_dir(slug: str, job_id: str) -> Path:
    d = _db.hermes_home() / "rolefit_artifacts" / slug / job_id
    d.mkdir(parents=True, exist_ok=True)
    return d


# --- CV (.docx) ------------------------------------------------------------

_CV_SYSTEM = (
    "You are an expert resume writer. Using the candidate's background and the target "
    "job, write a tailored, ATS-friendly resume. Emphasize the overlap with the job; "
    "never invent experience the candidate doesn't have. Respond ONLY as JSON: "
    '{"name":"","headline":"","summary":"","skills":["",...],'
    '"experience":[{"title":"","company":"","period":"","bullets":["",...]}],'
    '"education":["",...]}'
)


def _cv_content(person_ctx: str, job_ctx: str, slug: str) -> dict[str, Any]:
    user = f"CANDIDATE:\n{person_ctx}\n\nTARGET JOB:\n{job_ctx}"
    try:
        d = _llm.complete_json(_CV_SYSTEM, user)
        if isinstance(d, dict):
            d.setdefault("name", slug)
            return d
    except Exception:
        pass
    return {"name": slug, "headline": "", "summary": person_ctx[:400],
            "skills": [], "experience": [], "education": []}


def _render_cv_docx(content: dict[str, Any], path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    h = doc.add_heading(content.get("name") or "Candidate", level=0)
    if content.get("headline"):
        p = doc.add_paragraph()
        run = p.add_run(content["headline"])
        run.italic = True
        run.font.size = Pt(11)
    if content.get("summary"):
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(content["summary"])
    if content.get("skills"):
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(str(s) for s in content["skills"]))
    if content.get("experience"):
        doc.add_heading("Experience", level=1)
        for e in content["experience"]:
            line = " — ".join(x for x in [e.get("title"), e.get("company")] if x)
            ph = doc.add_paragraph()
            ph.add_run(line).bold = True
            if e.get("period"):
                ph.add_run(f"  ({e['period']})").italic = True
            for b in e.get("bullets") or []:
                doc.add_paragraph(str(b), style="List Bullet")
    if content.get("education"):
        doc.add_heading("Education", level=1)
        for ed in content["education"]:
            doc.add_paragraph(str(ed), style="List Bullet")
    doc.save(str(path))


# --- text artifacts --------------------------------------------------------

_COVER_SYSTEM = (
    "Write a concise, specific cover letter (~250 words) for this candidate and job. "
    "Lead with genuine fit; reference the company/role; no clichés or invented facts. "
    "Output the letter as Markdown."
)
_PREP_SYSTEM = (
    "You are an interview coach. Produce focused interview prep for this candidate and "
    "job as Markdown: likely technical + behavioral questions (with brief answer angles "
    "based on the candidate's background), and 3 smart questions to ask. Be specific."
)
_LEARN_SYSTEM = (
    "You are a career mentor. Given the candidate, the job, and the skills GAP, produce "
    "a short, actionable learning plan (Markdown): prioritized skills to close the gap, "
    "with concrete resources/steps and a rough timeline. Only address real gaps."
)
_OUTREACH_SYSTEM = (
    "Write a short, warm cold-outreach email to the hiring manager/recruiter for this "
    "job from this candidate (~120 words). Specific hook tied to the role/company, clear "
    "ask for a chat, no clichés. Output Markdown with a subject line."
)
_LINKEDIN_SYSTEM = (
    "Write a LinkedIn connection-request note (max 280 chars) from this candidate to "
    "someone at the company, referencing the role naturally. Output just the note text."
)
_KEYWORDS_SYSTEM = (
    "You are an ATS optimization assistant. From the job description, extract the exact "
    "keywords/skills an ATS will scan for, and mark which the candidate already has vs "
    "should add to their resume. Output Markdown: a 'Must-have keywords' table "
    "(keyword | candidate has? yes/no) + a short 'Add these to your resume' list."
)


def _text(system: str, person_ctx: str, job_ctx: str, extra: str = "") -> str:
    user = f"CANDIDATE:\n{person_ctx}\n\nJOB:\n{job_ctx}"
    if extra:
        user += f"\n\n{extra}"
    try:
        return _llm.complete(system, user, temperature=0.4).strip()
    except Exception as e:
        return f"_Generation failed: {e}_"


# --- orchestration ---------------------------------------------------------

_KINDS = ("cv", "cover", "interview", "learning", "outreach", "linkedin", "keywords")

_TEXT_GEN = {
    "cover": ("cover_letter.md", lambda s, pc, jc, gap: _text(_COVER_SYSTEM, pc, jc)),
    "interview": ("interview_prep.md", lambda s, pc, jc, gap: _text(_PREP_SYSTEM, pc, jc)),
    "learning": ("learning_plan.md", lambda s, pc, jc, gap: _text(_LEARN_SYSTEM, pc, jc, f"SKILLS GAP: {', '.join(gap) or 'none'}")),
    "outreach": ("outreach_email.md", lambda s, pc, jc, gap: _text(_OUTREACH_SYSTEM, pc, jc)),
    "linkedin": ("linkedin_message.md", lambda s, pc, jc, gap: _text(_LINKEDIN_SYSTEM, pc, jc)),
    "keywords": ("ats_keywords.md", lambda s, pc, jc, gap: _text(_KEYWORDS_SYSTEM, pc, jc)),
}
_PATH_COL = {
    "cv": "cv_path", "cover": "cover_path", "interview": "interview_path",
    "learning": "learning_path",
}


def generate(
    slug: str,
    job_id: str,
    *,
    tenant_id: str = _db.DEFAULT_TENANT,
    what: Optional[list[str]] = None,
) -> dict[str, Any]:
    """Generate the requested artifacts for one (person, job). Returns paths."""
    what = [w for w in (what or _KINDS) if w in _KINDS]
    person_ctx = _scoring.person_context(slug, tenant_id=tenant_id)
    job = _jobs.get_job(job_id, tenant_id=tenant_id)
    if not job:
        return {"error": "job not found"}
    job_ctx = _jobs.job_context_for_scoring(job_id, tenant_id=tenant_id)
    d = _art_dir(slug, job_id)

    conn = _db.connect()
    row = conn.execute(
        "SELECT gap_json FROM analyses WHERE tenant_id=? AND job_id=? AND person_id=?",
        (tenant_id, job_id, slug),
    ).fetchone()
    gap = json.loads(row["gap_json"]) if row and row["gap_json"] else []

    generated: list[str] = []
    db_paths: dict[str, str] = {}
    if "cv" in what:
        p = d / "cv.docx"
        _render_cv_docx(_cv_content(person_ctx, job_ctx, slug), p)
        generated.append("cv")
        db_paths["cv_path"] = str(p)
    for kind in what:
        if kind not in _TEXT_GEN:
            continue
        fname, fn = _TEXT_GEN[kind]
        p = d / fname
        p.write_text(fn(slug, person_ctx, job_ctx, gap), encoding="utf-8")
        generated.append(kind)
        if kind in _PATH_COL:
            db_paths[_PATH_COL[kind]] = str(p)

    # record on the analysis row (stage 2)
    now = time.time()
    sets = ", ".join(f"{k}=?" for k in db_paths)
    if sets:
        vals = list(db_paths.values()) + [now, tenant_id, job_id, slug]
        conn.execute(
            f"UPDATE analyses SET {sets}, stage=2, status='generated', updated_at=? "
            "WHERE tenant_id=? AND job_id=? AND person_id=?",
            vals,
        )
        conn.commit()
    return {"person": slug, "job_id": job_id, "generated": generated, **db_paths}


def generate_for_matches(
    slug: str,
    *,
    tenant_id: str = _db.DEFAULT_TENANT,
    min_score: int = 80,
    limit: int = 20,
    only_new: bool = True,
) -> dict[str, Any]:
    """Auto-generate for a person's strong matches (score ≥ min_score).

    only_new (default) skips matches already generated (stage>=2) — so the daily
    pipeline only does work for NEWLY scored strong matches, never re-generates.
    """
    matches = _scoring.list_matches(tenant_id=tenant_id, person=slug,
                                    min_score=min_score, limit=limit)
    if only_new:
        matches = [m for m in matches if (m.get("stage") or 1) < 2]
    done = [generate(slug, m["job_id"], tenant_id=tenant_id) for m in matches]
    return {"person": slug, "min_score": min_score, "count": len(done), "results": done}
